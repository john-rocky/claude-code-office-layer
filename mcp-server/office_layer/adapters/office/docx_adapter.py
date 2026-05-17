"""python-docx adapter — Word .docx."""

from __future__ import annotations

from pathlib import Path

from ...models import ChunkKind, DocumentChunk, DocumentKind, ExtractedTable, ExtractionMethod
from ..base import ExtractionResult


PARAGRAPHS_PER_CHUNK = 12


class PythonDocxAdapter:
    name = "python-docx"
    supported_kinds = (DocumentKind.DOCX,)

    def __init__(self) -> None:
        try:
            import docx  # noqa: F401
            self._available = True
        except ImportError:
            self._available = False

    def is_available(self) -> bool:
        return self._available

    def extract(
        self,
        path: Path,
        kind: DocumentKind,
        *,
        document_id: str,
        workspace_id: str,
    ) -> ExtractionResult:
        try:
            import docx
        except ImportError as exc:
            return ExtractionResult(error=f"python-docx missing: {exc}")

        try:
            d = docx.Document(str(path))
        except Exception as exc:
            return ExtractionResult(error=f"python-docx failed to open: {exc}")

        chunks: list[DocumentChunk] = []
        tables: list[ExtractedTable] = []
        ordinal = 0
        buffer: list[str] = []
        current_heading: str | None = None

        def flush() -> None:
            nonlocal ordinal, buffer
            if not buffer:
                return
            text = "\n".join(buffer).strip()
            if text:
                chunks.append(
                    DocumentChunk(
                        id=DocumentChunk.make_id(document_id, ordinal),
                        document_id=document_id,
                        workspace_id=workspace_id,
                        kind=ChunkKind.DOCX_SECTION,
                        ordinal=ordinal,
                        section_heading=current_heading,
                        text=text,
                        char_count=len(text),
                        confidence=1.0,
                        extraction_method=ExtractionMethod.DOCX_PARSE,
                    )
                )
                ordinal += 1
            buffer = []

        for para in d.paragraphs:
            style = (para.style.name or "").lower() if para.style else ""
            text = para.text.strip()
            if not text:
                continue
            if "heading" in style:
                flush()
                current_heading = text
                buffer.append(text)
                continue
            buffer.append(text)
            if len(buffer) >= PARAGRAPHS_PER_CHUNK:
                flush()
        flush()

        for tbl in d.tables:
            rows = []
            for row in tbl.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if not rows:
                continue
            headers = rows[0]
            body = rows[1:] if len(rows) > 1 else []
            tables.append(
                ExtractedTable(
                    document_id=document_id,
                    headers=headers,
                    rows=body,
                    confidence=0.95,
                )
            )

        return ExtractionResult(
            chunks=chunks,
            tables=tables,
            extraction_method=ExtractionMethod.DOCX_PARSE,
        )
